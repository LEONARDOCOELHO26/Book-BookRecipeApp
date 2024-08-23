import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter import ttk
import sqlite3
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from docx import Document
from docx.shared import Pt

# Função para selecionar o banco de dados
def selecionar_banco():
    filepath = filedialog.askopenfilename(filetypes=[("SQLite files", "*.db"), ("All files", "*.*")])
    if filepath:
        global banco_dados
        banco_dados = filepath
        setup_database()
        carregar_receitas()
    else:
        messagebox.showwarning("Atenção", "Nenhum banco de dados selecionado. O aplicativo será fechado.")
        root.destroy()

# Configuração do Banco de Dados
def setup_database():
    conn = sqlite3.connect(banco_dados)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS receitas
                      (id INTEGER PRIMARY KEY, titulo TEXT, ingredientes TEXT, modo_preparo TEXT)''')
    conn.commit()
    conn.close()

# Funções do Aplicativo
def adicionar_receita():
    titulo = titulo_entry.get()
    ingredientes = ingredientes_text.get("1.0", tk.END)
    modo_preparo = modo_preparo_text.get("1.0", tk.END)
    if titulo.strip() and ingredientes.strip() and modo_preparo.strip():
        try:
            conn = sqlite3.connect(banco_dados)
            cursor = conn.cursor()
            cursor.execute('INSERT INTO receitas (titulo, ingredientes, modo_preparo) VALUES (?, ?, ?)', (titulo, ingredientes, modo_preparo))
            conn.commit()
            conn.close()
            titulo_entry.delete(0, tk.END)
            ingredientes_text.delete("1.0", tk.END)
            modo_preparo_text.delete("1.0", tk.END)
            carregar_receitas()
            messagebox.showinfo("Sucesso", "Receita adicionada com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao adicionar a receita: {e}")
    else:
        messagebox.showwarning("Atenção", "Título, ingredientes e modo de preparo não podem ser vazios.")

def carregar_receitas():
    lista_receitas.delete(*lista_receitas.get_children())
    try:
        conn = sqlite3.connect(banco_dados)
        cursor = conn.cursor()
        cursor.execute('SELECT id, titulo FROM receitas')
        for row in cursor.fetchall():
            lista_receitas.insert("", "end", values=row)
        conn.close()
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao carregar as receitas: {e}")

def visualizar_receita(event):
    selecionado = lista_receitas.selection()
    if selecionado:
        id_receita = lista_receitas.item(selecionado)['values'][0]
        try:
            conn = sqlite3.connect(banco_dados)
            cursor = conn.cursor()
            cursor.execute('SELECT titulo, ingredientes, modo_preparo FROM receitas WHERE id = ?', (id_receita,))
            receita = cursor.fetchone()
            conn.close()
            titulo_entry.delete(0, tk.END)
            ingredientes_text.delete("1.0", tk.END)
            modo_preparo_text.delete("1.0", tk.END)
            titulo_entry.insert(0, receita[0])
            ingredientes_text.insert("1.0", receita[1])
            modo_preparo_text.insert("1.0", receita[2])
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao visualizar a receita: {e}")

def editar_receita():
    selecionado = lista_receitas.selection()
    if selecionado:
        id_receita = lista_receitas.item(selecionado)['values'][0]
        titulo = titulo_entry.get()
        ingredientes = ingredientes_text.get("1.0", tk.END)
        modo_preparo = modo_preparo_text.get("1.0", tk.END)
        if titulo.strip() and ingredientes.strip() and modo_preparo.strip():
            try:
                conn = sqlite3.connect(banco_dados)
                cursor = conn.cursor()
                cursor.execute('UPDATE receitas SET titulo = ?, ingredientes = ?, modo_preparo = ? WHERE id = ?', (titulo, ingredientes, modo_preparo, id_receita))
                conn.commit()
                conn.close()
                carregar_receitas()
                messagebox.showinfo("Sucesso", "Receita editada com sucesso!")
            except Exception as e:
                messagebox.showerror("Erro", f"Ocorreu um erro ao editar a receita: {e}")
        else:
            messagebox.showwarning("Atenção", "Título, ingredientes e modo de preparo não podem ser vazios.")
    else:
        messagebox.showwarning("Atenção", "Nenhuma receita selecionada para edição.")

def excluir_receita():
    selecionado = lista_receitas.selection()
    if selecionado:
        id_receita = lista_receitas.item(selecionado)['values'][0]
        try:
            conn = sqlite3.connect(banco_dados)
            cursor = conn.cursor()
            cursor.execute('DELETE FROM receitas WHERE id = ?', (id_receita,))
            conn.commit()
            conn.close()
            titulo_entry.delete(0, tk.END)
            ingredientes_text.delete("1.0", tk.END)
            modo_preparo_text.delete("1.0", tk.END)
            carregar_receitas()
            messagebox.showinfo("Sucesso", "Receita excluída com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao excluir a receita: {e}")
    else:
        messagebox.showwarning("Atenção", "Nenhuma receita selecionada para exclusão.")

def visualizar_receita_anterior():
    selecionado = lista_receitas.selection()
    if selecionado:
        item_atual = lista_receitas.item(selecionado)
        index_atual = lista_receitas.index(selecionado)
        if index_atual > 0:
            item_anterior = lista_receitas.get_children()[index_atual - 1]
            lista_receitas.selection_set(item_anterior)
            lista_receitas.event_generate("<<TreeviewSelect>>")

def extrair_livro_pdf():
    try:
        conn = sqlite3.connect(banco_dados)
        cursor = conn.cursor()
        cursor.execute('SELECT titulo, ingredientes, modo_preparo FROM receitas')
        receitas = cursor.fetchall()
        conn.close()

        if receitas:
            filepath = filedialog.asksaveasfilename(defaultextension=".pdf",
                                                    filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")])
            if filepath:
                doc = SimpleDocTemplate(filepath, pagesize=letter)
                styles = getSampleStyleSheet()
                elements = []

                for titulo, ingredientes, modo_preparo in receitas:
                    elements.append(Paragraph(titulo, styles['Title']))
                    elements.append(Spacer(1, 0.2 * inch))
                    elements.append(Paragraph("Ingredientes:<br />" + ingredientes.replace('\n', '<br />'), styles['BodyText']))
                    elements.append(Spacer(1, 0.2 * inch))
                    elements.append(Paragraph("Modo de Preparo:<br />" + modo_preparo.replace('\n', '<br />'), styles['BodyText']))
                    elements.append(PageBreak())

                doc.build(elements)
                messagebox.showinfo("Sucesso", "Livro de receitas extraído com sucesso em PDF!")
        else:
            messagebox.showwarning("Atenção", "Não há receitas para extrair.")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao extrair o livro: {e}")

def extrair_livro_word():
    try:
        conn = sqlite3.connect(banco_dados)
        cursor = conn.cursor()
        cursor.execute('SELECT titulo, ingredientes, modo_preparo FROM receitas')
        receitas = cursor.fetchall()
        conn.close()

        if receitas:
            filepath = filedialog.asksaveasfilename(defaultextension=".docx",
                                                    filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
            if filepath:
                doc = Document()
                for titulo, ingredientes, modo_preparo in receitas:
                    doc.add_heading(titulo, level=1)
                    doc.add_heading('Ingredientes', level=2)
                    p = doc.add_paragraph(ingredientes)
                    p.style.font.size = Pt(12)
                    doc.add_heading('Modo de Preparo', level=2)
                    p = doc.add_paragraph(modo_preparo)
                    p.style.font.size = Pt(12)
                    doc.add_page_break()

                doc.save(filepath)
                messagebox.showinfo("Sucesso", "Livro de receitas extraído com sucesso em Word!")
        else:
            messagebox.showwarning("Atenção", "Não há receitas para extrair.")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao extrair o livro: {e}")

# Configuração da Interface Tkinter
root = tk.Tk()
root.title("Editor de Livro de Receitas")

# Configurar a janela para tela cheia
root.state('zoomed')

# Uso de ttk para um design mais moderno
mainframe = ttk.Frame(root, padding="10 10 10 10")
mainframe.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

# Entrada para o título da receita
ttk.Label(mainframe, text="Título").grid(row=0, column=0, sticky=tk.W)
titulo_entry = ttk.Entry(mainframe, width=50)
titulo_entry.grid(row=1, column=0, sticky=(tk.W, tk.E))

# Entrada para os ingredientes da receita
ttk.Label(mainframe, text="Ingredientes").grid(row=2, column=0, sticky=tk.W)
ingredientes_text = tk.Text(mainframe, width=60, height=10, wrap='word')
ingredientes_text.grid(row=3, column=0, sticky=(tk.W, tk.E))

# Entrada para o modo de preparo da receita
ttk.Label(mainframe, text="Modo de Preparo").grid(row=4, column=0, sticky=tk.W)
modo_preparo_text = tk.Text(mainframe, width=60, height=10, wrap='word')
modo_preparo_text.grid(row=5, column=0, sticky=(tk.W, tk.E))

# Botões de ação
button_frame = ttk.Frame(mainframe)
button_frame.grid(row=6, column=0, sticky=tk.E)
ttk.Button(button_frame, text="Adicionar Receita", command=adicionar_receita).grid(row=0, column=0)
ttk.Button(button_frame, text="Editar Receita", command=editar_receita).grid(row=0, column=1)
ttk.Button(button_frame, text="Excluir Receita", command=excluir_receita).grid(row=0, column=2)
ttk.Button(button_frame, text="Receita Anterior", command=visualizar_receita_anterior).grid(row=0, column=3)
ttk.Button(button_frame, text="Extrair Livro (PDF)", command=extrair_livro_pdf).grid(row=0, column=4)
ttk.Button(button_frame, text="Extrair Livro (Word)", command=extrair_livro_word).grid(row=0, column=5)

# Lista de receitas
ttk.Label(mainframe, text="Receitas").grid(row=7, column=0, sticky=tk.W)
colunas = ("ID", "Título")
lista_receitas = ttk.Treeview(mainframe, columns=colunas, show='headings')
for col in colunas:
    lista_receitas.heading(col, text=col)
lista_receitas.grid(row=8, column=0, sticky=(tk.W, tk.E))
lista_receitas.bind("<<TreeviewSelect>>", visualizar_receita)

# Configurações de redimensionamento
root.columnconfigure(0, weight=1)
root.rowconfigure(0, weight=1)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(3, weight=1)
mainframe.rowconfigure(6, weight=1)

# Selecionar banco de dados ao iniciar
root.after(100, selecionar_banco)

# Executar a aplicação Tkinter
root.mainloop()
