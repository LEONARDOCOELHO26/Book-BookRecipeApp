import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter import ttk
import sqlite3
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from docx import Document
from docx.shared import Pt

# Função para selecionar o banco de dados
def selecionar_banco():
    filepath = filedialog.askopenfilename(filetypes=[("SQLite files", "*.db"), ("All files", "*.*")])
    if filepath:
        global banco_dados
        banco_dados = filepath
        setup_database()
        carregar_capitulos()
    else:
        messagebox.showwarning("Atenção", "Nenhum banco de dados selecionado. O aplicativo será fechado.")
        root.destroy()

# Configuração do Banco de Dados
def setup_database():
    conn = sqlite3.connect(banco_dados)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS capitulos
                      (id INTEGER PRIMARY KEY, titulo TEXT, conteudo TEXT)''')
    conn.commit()
    conn.close()

# Funções do Aplicativo
def adicionar_capitulo():
    titulo = titulo_entry.get()
    conteudo = conteudo_text.get("1.0", tk.END)
    if titulo.strip() and conteudo.strip():
        try:
            conn = sqlite3.connect(banco_dados)
            cursor = conn.cursor()
            cursor.execute('INSERT INTO capitulos (titulo, conteudo) VALUES (?, ?)', (titulo, conteudo))
            conn.commit()
            conn.close()
            titulo_entry.delete(0, tk.END)
            conteudo_text.delete("1.0", tk.END)
            carregar_capitulos()
            messagebox.showinfo("Sucesso", "Capítulo adicionado com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao adicionar o capítulo: {e}")
    else:
        messagebox.showwarning("Atenção", "Título e conteúdo não podem ser vazios.")

def carregar_capitulos():
    lista_capitulos.delete(*lista_capitulos.get_children())
    try:
        conn = sqlite3.connect(banco_dados)
        cursor = conn.cursor()
        cursor.execute('SELECT id, titulo FROM capitulos')
        for row in cursor.fetchall():
            lista_capitulos.insert("", "end", values=row)
        conn.close()
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao carregar os capítulos: {e}")

def visualizar_capitulo(event):
    selecionado = lista_capitulos.selection()
    if selecionado:
        id_capitulo = lista_capitulos.item(selecionado)['values'][0]
        try:
            conn = sqlite3.connect(banco_dados)
            cursor = conn.cursor()
            cursor.execute('SELECT titulo, conteudo FROM capitulos WHERE id = ?', (id_capitulo,))
            capitulo = cursor.fetchone()
            conn.close()
            titulo_entry.delete(0, tk.END)
            conteudo_text.delete("1.0", tk.END)
            titulo_entry.insert(0, capitulo[0])
            conteudo_text.insert("1.0", capitulo[1])
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao visualizar o capítulo: {e}")

def editar_capitulo():
    selecionado = lista_capitulos.selection()
    if selecionado:
        id_capitulo = lista_capitulos.item(selecionado)['values'][0]
        titulo = titulo_entry.get()
        conteudo = conteudo_text.get("1.0", tk.END)
        if titulo.strip() and conteudo.strip():
            try:
                conn = sqlite3.connect(banco_dados)
                cursor = conn.cursor()
                cursor.execute('UPDATE capitulos SET titulo = ?, conteudo = ? WHERE id = ?', (titulo, conteudo, id_capitulo))
                conn.commit()
                conn.close()
                carregar_capitulos()
                messagebox.showinfo("Sucesso", "Capítulo editado com sucesso!")
            except Exception as e:
                messagebox.showerror("Erro", f"Ocorreu um erro ao editar o capítulo: {e}")
        else:
            messagebox.showwarning("Atenção", "Título e conteúdo não podem ser vazios.")
    else:
        messagebox.showwarning("Atenção", "Nenhum capítulo selecionado para edição.")

def excluir_capitulo():
    selecionado = lista_capitulos.selection()
    if selecionado:
        id_capitulo = lista_capitulos.item(selecionado)['values'][0]
        try:
            conn = sqlite3.connect(banco_dados)
            cursor = conn.cursor()
            cursor.execute('DELETE FROM capitulos WHERE id = ?', (id_capitulo,))
            conn.commit()
            conn.close()
            titulo_entry.delete(0, tk.END)
            conteudo_text.delete("1.0", tk.END)
            carregar_capitulos()
            messagebox.showinfo("Sucesso", "Capítulo excluído com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao excluir o capítulo: {e}")
    else:
        messagebox.showwarning("Atenção", "Nenhum capítulo selecionado para exclusão.")

def visualizar_capitulo_anterior():
    selecionado = lista_capitulos.selection()
    if selecionado:
        item_atual = lista_capitulos.item(selecionado)
        index_atual = lista_capitulos.index(selecionado)
        if index_atual > 0:
            item_anterior = lista_capitulos.get_children()[index_atual - 1]
            lista_capitulos.selection_set(item_anterior)
            lista_capitulos.event_generate("<<TreeviewSelect>>")

def extrair_livro_pdf():
    try:
        conn = sqlite3.connect(banco_dados)
        cursor = conn.cursor()
        cursor.execute('SELECT titulo, conteudo FROM capitulos')
        capitulos = cursor.fetchall()
        conn.close()

        if capitulos:
            filepath = filedialog.asksaveasfilename(defaultextension=".pdf",
                                                    filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")])
            if filepath:
                doc = SimpleDocTemplate(filepath, pagesize=letter)
                styles = getSampleStyleSheet()
                elements = []

                for titulo, conteudo in capitulos:
                    elements.append(Paragraph(titulo, styles['Title']))
                    elements.append(Spacer(1, 0.2 * inch))
                    elements.append(Paragraph(conteudo.replace('\n', '<br />'), styles['BodyText']))
                    elements.append(PageBreak())

                doc.build(elements)
                messagebox.showinfo("Sucesso", "Livro extraído com sucesso em PDF!")
        else:
            messagebox.showwarning("Atenção", "Não há capítulos para extrair.")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao extrair o livro: {e}")

def extrair_livro_word():
    try:
        conn = sqlite3.connect(banco_dados)
        cursor = conn.cursor()
        cursor.execute('SELECT titulo, conteudo FROM capitulos')
        capitulos = cursor.fetchall()
        conn.close()

        if capitulos:
            filepath = filedialog.asksaveasfilename(defaultextension=".docx",
                                                    filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
            if filepath:
                doc = Document()
                for titulo, conteudo in capitulos:
                    doc.add_heading(titulo, level=1)
                    p = doc.add_paragraph(conteudo)
                    p.style.font.size = Pt(12)
                    doc.add_page_break()

                doc.save(filepath)
                messagebox.showinfo("Sucesso", "Livro extraído com sucesso em Word!")
        else:
            messagebox.showwarning("Atenção", "Não há capítulos para extrair.")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao extrair o livro: {e}")

# Configuração da Interface Tkinter
root = tk.Tk()
root.title("Editor de Livro")

# Configurar a janela para tela cheia
root.state('zoomed')

# Uso de ttk para um design mais moderno
mainframe = ttk.Frame(root, padding="10 10 10 10")
mainframe.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

# Entrada para o título do capítulo
ttk.Label(mainframe, text="Título").grid(row=0, column=0, sticky=tk.W)
titulo_entry = ttk.Entry(mainframe, width=50)
titulo_entry.grid(row=1, column=0, sticky=(tk.W, tk.E))

# Entrada para o conteúdo do capítulo
ttk.Label(mainframe, text="Conteúdo").grid(row=2, column=0, sticky=tk.W)
conteudo_text = tk.Text(mainframe, width=60, height=15, wrap='word')
conteudo_text.grid(row=3, column=0, sticky=(tk.W, tk.E))

# Botões de ação
button_frame = ttk.Frame(mainframe)
button_frame.grid(row=4, column=0, sticky=tk.E)
ttk.Button(button_frame, text="Adicionar Capítulo", command=adicionar_capitulo).grid(row=0, column=0)
ttk.Button(button_frame, text="Editar Capítulo", command=editar_capitulo).grid(row=0, column=1)
ttk.Button(button_frame, text="Excluir Capítulo", command=excluir_capitulo).grid(row=0, column=2)
ttk.Button(button_frame, text="Capítulo Anterior", command=visualizar_capitulo_anterior).grid(row=0, column=3)
ttk.Button(button_frame, text="Extrair Livro (PDF)", command=extrair_livro_pdf).grid(row=0, column=4)
ttk.Button(button_frame, text="Extrair Livro (Word)", command=extrair_livro_word).grid(row=0, column=5)

# Lista de capítulos
ttk.Label(mainframe, text="Capítulos").grid(row=5, column=0, sticky=tk.W)
colunas = ("ID", "Título")
lista_capitulos = ttk.Treeview(mainframe, columns=colunas, show='headings')
for col in colunas:
    lista_capitulos.heading(col, text=col)
lista_capitulos.grid(row=6, column=0, sticky=(tk.W, tk.E))
lista_capitulos.bind("<<TreeviewSelect>>", visualizar_capitulo)

# Configurações de redimensionamento
root.columnconfigure(0, weight=1)
root.rowconfigure(0, weight=1)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(3, weight=1)
mainframe.rowconfigure(6, weight=1)

# Selecionar banco de dados ao iniciar
root.after(100, selecionar_banco)

# Executar a aplicação Tkinter
root.mainloop()
